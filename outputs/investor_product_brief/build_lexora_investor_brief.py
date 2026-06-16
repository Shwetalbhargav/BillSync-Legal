from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "outputs/investor_product_brief/Lexora_Investor_Product_Brief.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALL_OUT = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", size=11, color=None, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALL_OUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(label + " ")
    set_font(r1, size=11, color=NAVY, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=11, color=None)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def fill_table(table, headers, rows, widths):
    set_table_width(table, widths)
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_font(r, size=10, color=NAVY, bold=True)
    for row_idx, values in enumerate(rows, start=1):
        row = table.rows[row_idx]
        for i, value in enumerate(values):
            cell = row.cells[i]
            set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            r = p.add_run(value)
            set_font(r, size=9.5)


def set_doc_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Lexora Investor Product Brief")
    set_font(r, size=9, color=GRAY)


def build():
    doc = Document()
    set_doc_styles(doc)
    add_footer(doc.sections[0])

    add_para(doc, "Investor Product Brief", size=11, color=GRAY, bold=True, after=2)
    title = add_para(doc, "Lexora", size=30, color=NAVY, bold=True, after=2)
    title.paragraph_format.keep_with_next = True
    add_para(
        doc,
        "AI-powered Legal Practice Management SaaS for Indian law firms",
        size=14,
        color=GRAY,
        after=14,
    )
    meta = doc.add_table(rows=5, cols=2)
    fill_table(
        meta,
        ["Investor Lens", "Product Position"],
        [
            ["Category", "Legal Practice Management SaaS: CRM + ERP + billing automation + AI"],
            ["Primary users", "Law firm owners, partners, lawyers, associates, finance teams, and admins"],
            ["Core promise", "Capture more billable work, invoice faster, reduce admin effort, and improve firm visibility"],
            ["India fit", "Matter workflows, court sync, GST/TDS, receivables, payroll, and operational reporting"],
        ],
        [2100, 7260],
    )

    add_heading(doc, "Executive Summary", 1)
    add_callout(
        doc,
        "Investment thesis:",
        "Lexora turns a law firm's scattered daily work into structured, billable, auditable, and AI-assisted operations.",
    )
    add_para(
        doc,
        "Most Indian law firms still run on a fragmented mix of Excel sheets, email threads, manual timesheets, court diaries, WhatsApp updates, and disconnected billing tools. This creates revenue leakage, slow invoice cycles, weak visibility, and avoidable partner/admin workload.",
    )
    add_para(
        doc,
        "Lexora is designed as a single operating system for a law firm. It connects client and matter records, task management, automated work capture, billing approvals, invoices, payments, finance, HR, compliance, documents, integrations, and AI assistance.",
    )
    add_para(
        doc,
        "The product is especially compelling because it combines practice management with productivity capture. That gives firms a direct financial reason to adopt: more accurate billable time, faster billing, stronger utilization insight, and less non-billable administrative effort.",
    )

    add_heading(doc, "The Problem Lexora Solves", 1)
    for item in [
        "Billable work is lost because lawyers forget to record emails, calls, research, drafting, court work, and follow-ups.",
        "Partners lack real-time visibility into matter progress, team workload, approval bottlenecks, and revenue leakage.",
        "Finance teams spend too much time reconciling time entries, invoices, payments, GST/TDS, and receivables manually.",
        "Lawyers spend expensive hours on low-value coordination, status updates, document searching, and administrative reporting.",
        "Existing tools are either generic practice-management systems or enterprise legal platforms, but few combine India-ready legal operations with work capture and AI.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Product Architecture", 1)
    arch = doc.add_table(rows=1 + 4, cols=3)
    fill_table(
        arch,
        ["Layer", "What It Does", "Strategic Value"],
        [
            ["Web app", "Central workspace for matters, billing, finance, HR, AI, documents, settings, and dashboards.", "Daily system of record for the firm."],
            ["Backend API", "Node/MongoDB service layer for users, clients, cases, tasks, work sessions, invoices, payments, analytics, integrations, and AI.", "Scalable product foundation with modular services."],
            ["Desktop agent", "Privacy-safe Windows-first activity tracker that records activity counts, active/inactive time, app name, and window title.", "Differentiated work capture and productivity intelligence."],
            ["Browser extension", "Captures browser-based work and supports setup/status flows for lawyers working inside web tools.", "Improves capture accuracy at the point of work."],
        ],
        [1600, 4600, 3160],
    )

    add_heading(doc, "Modules Available", 1)
    modules = [
        ["Dashboard and profile", "Role-aware workspace, daily summaries, notifications, search, setup status.", "Gives every user a clear command center.", "Improves adoption and daily engagement."],
        ["Clients and CRM", "Client list, profiles, contacts, billing summaries, communications.", "Improves relationship management and client history.", "Supports retention and cross-sell visibility."],
        ["Matters and cases", "Matter list, creation, details, timeline, team, documents, audit trail, assignments.", "Creates a single source of truth for legal work.", "Reduces delivery risk and partner dependency."],
        ["Tasks and daily work", "My tasks, team board, task details, My Work Today.", "Reduces coordination overhead and missed deadlines.", "Raises execution discipline across teams."],
        ["Work Meter and time capture", "Work sessions, manual time, captured work, submit work, approval queue.", "Converts activity into reviewable billable work.", "Directly reduces revenue leakage."],
        ["Billing and invoices", "Billables, approvals, rate cards, invoice builder, invoice detail, invoice lines.", "Shortens quote-to-cash and reduces billing leakage.", "Improves realization and billing throughput."],
        ["Payments and finance", "Payments, reconciliation, receivables aging, revenue, KPIs, analytics, reports.", "Improves cash visibility and partner decision-making.", "Strengthens cash flow and forecasting."],
        ["India tax and compliance", "GST dashboard, TDS management, compliance settings, audit logs.", "Supports India-specific finance and governance needs.", "Creates local-market defensibility."],
        ["People and HR", "HR dashboard, people directory, attendance, leave, payroll, compensation, workload analytics.", "Connects people cost, utilization, and productivity.", "Expands ACV beyond legal workflow."],
        ["Court and calendar", "Hearings, manual court time, court sync, case match, verdict details.", "Reduces manual court tracking and missed update risk.", "Improves deadline reliability."],
        ["Documents and storage", "Document storage, upload, viewer, storage settings, cloud providers.", "Keeps legal artifacts attached to matters.", "Improves knowledge reuse and audit readiness."],
        ["Communications", "Communication center, WhatsApp/SMS/email logs, client interactions.", "Preserves context and reduces scattered conversations.", "Improves service quality and accountability."],
        ["AI assistant", "Matter summaries, document summaries, Q&A, drafting, research, email writer, product coach.", "Saves lawyer time and improves knowledge reuse.", "Creates premium upsell and usage depth."],
        ["Integrations", "Zoho, cloud storage, integration logs, future legal/accounting connectors.", "Fits into existing firm operations.", "Reduces switching friction."],
        ["Admin and security", "User management, roles, firms, profiles, security, permissions.", "Enables multi-role firm-wide deployment.", "Supports larger firm sales."],
    ]
    table = doc.add_table(rows=1 + len(modules), cols=4)
    fill_table(
        table,
        ["Module", "Capabilities", "User Benefit", "Business Impact"],
        [[m[0], m[1], m[2], m[3]] for m in modules],
        [1800, 3300, 2500, 1760],
    )

    add_heading(doc, "How Lexora Increases Productivity", 1)
    add_para(doc, "Lexora improves productivity through four compounding mechanisms:")
    for item in [
        "Capture: Lawyers spend less time remembering and reconstructing work because sessions, activities, browser work, and manual entries are centralized.",
        "Review: Captured work becomes structured, editable, and approvable, so partners can control quality without chasing every user manually.",
        "Automation: Billing, invoices, payments, reminders, dashboards, court updates, and reports reduce repetitive administrative work.",
        "Intelligence: AI summaries, document Q&A, research assistance, and matter insights reduce time spent searching, drafting, and preparing updates.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Value for Money, Effort, and Knowledge", 1)
    value = doc.add_table(rows=1 + 4, cols=4)
    fill_table(
        value,
        ["Value Lever", "Before Lexora", "With Lexora", "Why It Matters"],
        [
            ["Money", "Missed billable time, delayed invoices, weak receivables control.", "Higher capture accuracy, faster approvals, stronger payment visibility.", "Even small increases in recovered billable hours can pay for the software."],
            ["Effort", "Manual timesheets, Excel reports, repeated follow-ups, duplicated data entry.", "Automated capture, centralized records, dashboards, and workflow queues.", "Admin and lawyer time moves from clerical work to client work."],
            ["Knowledge", "Matter context lives in emails, documents, chats, and individual memory.", "Matter timelines, documents, AI summaries, audit trails, and searchable records.", "Firm knowledge becomes reusable and less dependent on individuals."],
            ["Control", "Partners see problems late, after billing or deadlines slip.", "Live workload, approval, revenue, and compliance visibility.", "Management can intervene earlier and make better operating decisions."],
        ],
        [1700, 2550, 2550, 2560],
    )

    add_heading(doc, "Customer ROI Narrative", 1)
    add_callout(
        doc,
        "Simple ROI story:",
        "If a firm recovers only a few additional billable hours per lawyer each month, Lexora can justify its subscription while also reducing admin burden.",
    )
    add_para(
        doc,
        "For customers, the product should be sold around measurable operational outcomes: fewer missed billables, faster invoice cycles, reduced partner review time, cleaner collections, better workload planning, and stronger compliance visibility.",
    )
    for item in [
        "Revenue lift from recovered time and fewer unbilled activities.",
        "Cash-flow improvement from faster invoice creation and receivables follow-up.",
        "Cost reduction from less manual coordination, spreadsheet reporting, and finance cleanup.",
        "Quality improvement from better matter records, document context, approval trails, and audit logs.",
        "Capacity improvement because lawyers and partners spend more time on paid legal work and less time on operations.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Positioning Against Competitors", 1)
    comp = doc.add_table(rows=1 + 5, cols=4)
    fill_table(
        comp,
        ["Product Type", "Typical Strength", "Typical Gap", "Lexora Advantage"],
        [
            ["Basic case software", "Matters, clients, court dates.", "Limited finance, HR, AI, and work capture.", "Broader firm operating system."],
            ["Global tools like Clio/MyCase", "Mature legal practice workflows.", "Expensive for India and less India-specific.", "India-ready pricing, GST/TDS, court sync, local workflows."],
            ["Enterprise legal platforms", "Corporate legal, contracts, litigation portfolios.", "Less optimized for daily law-firm billing productivity.", "Built around law firm revenue and utilization."],
            ["Time trackers", "Activity and productivity measurement.", "Not legal-native and not connected to billing/matters.", "Legal time capture tied to matters, approvals, and invoices."],
            ["AI legal tools", "Drafting, summaries, research.", "Often standalone without operations workflow.", "AI embedded inside the firm workflow."],
        ],
        [2000, 2500, 2450, 2410],
    )

    add_heading(doc, "Suggested SaaS Packaging", 1)
    pricing = doc.add_table(rows=1 + 4, cols=4)
    fill_table(
        pricing,
        ["Plan", "Target Customer", "Indicative Price", "Primary Value"],
        [
            ["Starter", "Solo lawyers and very small firms", "Rs 999/user/month", "Matters, tasks, clients, basic time and billing."],
            ["Professional", "Growing firms", "Rs 2,499/user/month", "Work capture, approvals, invoices, dashboards, AI basics."],
            ["Firm Plus", "Established firms", "Rs 4,999/user/month", "Finance, HR, GST/TDS, advanced reporting, integrations."],
            ["Enterprise", "Large firms and multi-office practices", "Custom, Rs 75,000+/month", "Migration, custom workflows, priority support, private deployment options."],
        ],
        [1550, 2950, 2100, 2760],
    )
    add_para(
        doc,
        "Pricing should remain below premium global tools while still communicating that Lexora is a productivity and revenue-recovery platform, not a simple case diary.",
        size=10.5,
        color=GRAY,
        italic=True,
    )

    add_heading(doc, "Why This Can Become a Valuable Company", 1)
    for item in [
        "Large, under-digitized Indian legal services market with persistent operational pain.",
        "Clear economic buyer: partners and firm owners who directly benefit from higher realization and faster collections.",
        "Multi-module expansion path creates account growth after initial adoption.",
        "Work capture and billing data can become a defensible intelligence layer for firm productivity and profitability.",
        "AI features become more valuable when connected to matter context, documents, billing, and firm workflows.",
        "Desktop and browser capture create product differentiation beyond standard CRM/case management systems.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Roadmap Priorities", 1)
    roadmap = doc.add_table(rows=1 + 4, cols=3)
    fill_table(
        roadmap,
        ["Priority", "Why It Matters", "Investor Relevance"],
        [
            ["Nail the billing loop", "Capture work, review it, approve it, invoice it, reconcile payment.", "Creates measurable ROI and strong retention."],
            ["Deepen India workflows", "Court sync, GST/TDS, local invoice formats, collections, WhatsApp client workflows.", "Improves local defensibility."],
            ["Strengthen AI inside matters", "Summaries, Q&A, drafting, and knowledge reuse from real matter context.", "Turns AI into workflow advantage, not a standalone feature."],
            ["Build analytics moat", "Utilization, realization, leakage, workload, profitability, and team performance.", "Creates executive dependence and upsell potential."],
        ],
        [2200, 4300, 2860],
    )

    add_heading(doc, "Investor Takeaway", 1)
    add_callout(
        doc,
        "Bottom line:",
        "Lexora has the potential to become the operating system for Indian law firms by combining CRM, ERP, billing automation, AI, and privacy-aware productivity capture in one platform.",
    )
    add_para(
        doc,
        "The most investable angle is the direct connection between product usage and customer economics. Lexora does not only organize work; it helps firms recover revenue, reduce administrative effort, preserve knowledge, and manage professional capacity with better visibility.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
