from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "outputs/Lexora_App_Overview_Guidde_Video.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 45)
MUTED = RGBColor(92, 104, 116)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "B8C7D9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Lexora overview | Guidde demo support")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Lexora App Overview")
    set_run_font(r, size=24, color=RGBColor(0, 0, 0), bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Guidde video reference for installation, workflow demonstration, and product explanation")
    set_run_font(r, size=12, color=MUTED)

    meta = [
        ("Project", "LexoraV1 legal productivity and billing platform"),
        ("Audience", "Hackathon judges, evaluators, demo viewers, and first-time users"),
        ("Goal", "Show how Lexora captures legal work, organizes it by client and matter, and turns it into billable records."),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    for row, (label, value) in zip(table.rows, meta):
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(5.1)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r, size=11, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r, size=11, color=INK)


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run_font(r, size=11, color=INK, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_run_font(r, size=11, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=INK)


def add_module_table(doc):
    modules = [
        ("Web dashboard", "Main user interface for clients, matters, tasks, time entries, billables, invoices, analytics, settings, and approvals."),
        ("Backend API", "Express and MongoDB service layer that manages authentication, firm data, legal records, billing entities, activity capture, and integrations."),
        ("Desktop agent", "Electron-based Windows agent that signs in users, tracks live work sessions, monitors app activity, and syncs captured activity to the backend."),
        ("Browser extension", "Chrome extension for Gmail typing-time capture, research capture, client/matter mapping, local retry queueing, and privacy-aware metadata submission."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    header = table.rows[0].cells
    header[0].width = Inches(1.875)
    header[1].width = Inches(4.625)
    for cell, text in zip(header, ("Component", "Role in the app")):
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    for component, role in modules:
        row = table.add_row().cells
        row[0].width = Inches(1.875)
        row[1].width = Inches(4.625)
        for cell in row:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = row[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(component)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
        p = row[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(role)
        set_run_font(r, size=10.5, color=INK)


def add_demo_flow_table(doc):
    rows = [
        ("1", "Sign in", "Show login and explain that Lexora connects legal work to a specific firm user."),
        ("2", "Create or select matter", "Open clients and matters to show how legal work is organized by client, matter, team, and billing context."),
        ("3", "Capture work", "Demonstrate manual time entry, desktop work meter, Gmail capture, or research capture."),
        ("4", "Review billables", "Show captured activity moving into time entries or billable records for review and approval."),
        ("5", "Generate invoice", "Open invoice builder and explain how approved billables become client-ready invoices."),
        ("6", "View analytics", "Show dashboards for revenue, workforce activity, time, attendance, and billing visibility."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)
    widths = [0.45, 1.55, 4.5]
    for cell, text, width in zip(table.rows[0].cells, ("Step", "Screen / action", "Narration point"), widths):
        cell.width = Inches(width)
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(text)
        set_run_font(r, size=10, color=DARK_BLUE, bold=True)
    for num, action, point in rows:
        cells = table.add_row().cells
        for cell, width in zip(cells, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cells[0].paragraphs[0].add_run(num)
        set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
        r = cells[1].paragraphs[0].add_run(action)
        set_run_font(r, size=10.5, color=INK, bold=True)
        r = cells[2].paragraphs[0].add_run(point)
        set_run_font(r, size=10.5, color=INK)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table, color="D4DFEC", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_run_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)


def build():
    doc = Document()
    style_document(doc)
    add_title_block(doc)

    doc.add_heading("One-Line Overview", level=1)
    add_callout(
        doc,
        "Demo message",
        "Lexora is a connected legal productivity and billing platform that captures work from web, desktop, and browser activity, maps it to clients and matters, and helps firms convert that work into reviewed billables and invoices.",
    )

    doc.add_heading("What the App Does", level=1)
    add_para(
        doc,
        "Lexora helps legal teams reduce missed billable time by connecting daily work activity to the business records a law firm already cares about: clients, matters, tasks, documents, time entries, billables, invoices, payments, and analytics.",
    )
    add_bullets(
        doc,
        [
            "Organizes legal work around clients, matters, assignments, documents, and tasks.",
            "Captures billable activity through manual time entry, desktop work sessions, Gmail typing activity, and research capture.",
            "Supports review workflows so captured activity can become accurate time entries and billables.",
            "Provides finance, invoice, payment, workforce, and productivity dashboards for firm visibility.",
            "Includes AI-assisted legal workflows for matter questions, document summaries, email drafting, and research support.",
        ],
    )

    doc.add_heading("Main Components", level=1)
    add_module_table(doc)

    doc.add_heading("How Lexora Works", level=1)
    add_numbered(
        doc,
        [
            "A user signs in through the web app or desktop agent.",
            "The firm creates clients and matters so each piece of work has legal and billing context.",
            "Work is captured manually or automatically from desktop activity, Gmail composition, and browser research.",
            "Captured activity is reviewed, mapped to the correct matter, and converted into time entries or billables.",
            "Approved billables can be used to build invoices and track payment or receivables status.",
            "Dashboards summarize billing, productivity, workforce activity, and firm operations.",
        ],
    )

    doc.add_heading("Recommended Guidde Demo Flow", level=1)
    add_demo_flow_table(doc)

    doc.add_heading("Installation Overview for the Video", level=1)
    add_para(doc, "Use this section when recording the installation part of the Guidde video.")
    add_bullets(
        doc,
        [
            "Start the backend API from apps/api-Lexora using npm install and npm run dev.",
            "Start the web dashboard from apps/web-Laxora using npm install and npm run dev.",
            "Start the desktop agent from apps/desktop-Lexora using npm install and npm run dev.",
            "Load the browser extension from apps/extension through the browser extensions page in Developer Mode.",
            "Confirm that the API, web app, desktop agent, and extension are all connected before recording the workflow demo.",
        ],
    )

    doc.add_heading("Features to Highlight", level=1)
    doc.add_heading("For lawyers and associates", level=2)
    add_bullets(
        doc,
        [
            "Track work sessions and manual time entries.",
            "Capture Gmail drafting time and browser research.",
            "Review captured activity before it becomes billable.",
            "Use AI tools for matter questions, document summaries, and drafting support.",
        ],
    )
    doc.add_heading("For partners and administrators", level=2)
    add_bullets(
        doc,
        [
            "Manage clients, matters, teams, tasks, permissions, and firm settings.",
            "Approve billables and create invoices from approved legal work.",
            "Monitor finance, payments, receivables, workforce, attendance, and productivity analytics.",
            "Use audit and compliance views to understand activity history and system state.",
        ],
    )

    doc.add_heading("Privacy and Reliability Points", level=1)
    add_bullets(
        doc,
        [
            "The Gmail extension focuses on metadata such as recipient, subject, user identity, timing, source references, and mapping data rather than sending email body text.",
            "Captured browser activity can be queued locally and retried if the backend is offline or a request fails.",
            "Client and matter mapping keeps captured work tied to the correct legal context before billing.",
            "Role-based access and protected routes keep users focused on the screens they are allowed to use.",
        ],
    )

    doc.add_heading("Suggested Closing Narration", level=1)
    add_para(
        doc,
        "Lexora brings legal operations, activity capture, billing review, and invoicing into one connected workflow. The result is less missed time, clearer matter-level visibility, and a smoother path from daily legal work to accurate client billing.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
