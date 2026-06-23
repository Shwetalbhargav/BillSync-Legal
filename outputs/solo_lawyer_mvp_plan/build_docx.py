import json, os, sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"E:\LexoraV1")
OUT = ROOT / "outputs" / "solo_lawyer_mvp_plan"
DATA = json.loads((OUT / "product_plan_data.json").read_text(encoding="utf-8"))
SKILL = Path(r"C:\Users\Bhargav Shah\.codex\plugins\cache\openai-primary-runtime\documents\26.614.11602\skills\documents")
sys.path.insert(0, str(SKILL / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "117A8B"
LIGHT_BLUE = "E8EEF5"
LIGHTER = "F4F6F9"
GRAY = "667085"
GRID = "D0D5DD"
GREEN = "147D64"
GOLD = "9A6700"
RED = "B42318"
WHITE = "FFFFFF"
INK = "182230"

def set_font(run, name="Calibri", size=11, bold=False, color=INK, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_border(cell, color=GRID, size="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:color"), color)

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)

def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True

def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("Page ")
    set_font(r, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r)
    return p

def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r)
    return p

def add_manual_number(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(f"{number}.  {text}")
    set_font(r)
    return p

def add_callout(doc, title, text, fill=LIGHTER, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0,0)
    shade(cell, fill)
    set_cell_border(cell, accent, "8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    set_font(r, size=9, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r = p2.add_run(text)
    set_font(r, size=11, bold=True, color=NAVY)
    apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_table(doc, headers, rows, weights, font_size=9.2, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(c, header_fill)
        set_cell_border(c, GRID)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(h))
        set_font(r, size=9, bold=True, color=WHITE)
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            c = cells[i]
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(c, WHITE if ri % 2 else "F8FAFC")
            set_cell_border(c, GRID)
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            r = p.add_run("" if value is None else str(value))
            set_font(r, size=font_size, color=INK)
    widths = column_widths_from_weights(weights, 9360)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in [
    ("Title", 30, NAVY, 0, 8), ("Subtitle", 13, GRAY, 0, 18),
    ("Heading 1", 16, BLUE, 18, 10), ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 11.5, DARK_BLUE, 10, 5)]:
    s = styles[name]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(size)
    s.font.bold = name != "Subtitle"
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

header = sec.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp.add_run("LEXORA SOLO  |  PRODUCT DELIVERY PLAN")
set_font(r, size=8.5, bold=True, color=GRAY)
footer = sec.footer
fp = footer.paragraphs[0]
add_page_field(fp)

# Cover / memo masthead
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
r = p.add_run("PRODUCT REQUIREMENTS + SCRUM DELIVERY PACK")
set_font(r, size=9.5, bold=True, color=TEAL)
p = doc.add_paragraph(style="Title")
p.add_run("Lexora Solo-Lawyer MVP")
p = doc.add_paragraph(style="Subtitle")
p.add_run("Eight-week implementation plan for a focused legal work-to-payment platform")

add_table(doc, ["Plan", "Decision"], [
    ["Target user", DATA["product"]["target"]],
    ["Core promise", DATA["product"]["positioning"]],
    ["Delivery model", DATA["product"]["duration"]],
    ["Release target", DATA["product"]["release"]],
    ["Repository", r"E:\LexoraV1"],
    ["Integration branch", "codex/solo-lawyer-mvp"]
], [1.65, 4.85], font_size=10)
add_callout(doc, "Product decision", "Position Lexora as a solo-lawyer work-to-payment product, not a legal ERP. Keep one owner per workspace, remove team approvals from the experience, and retain strict workspace isolation between customers.")

doc.add_heading("1. Product strategy", level=1)
doc.add_heading("Problem", level=2)
doc.add_paragraph("Independent lawyers lose billable time across email, research, calls, drafting and court work; reconstructing that work later delays invoicing and obscures collections. Lexora should make capture and billing a single, low-friction daily habit.")
doc.add_heading("North-star workflow", level=2)
add_callout(doc, "Workflow", DATA["product"]["workflow"], fill="E9F7F5", accent=GREEN)
doc.add_heading("Product objectives and measures", level=2)
add_table(doc, ["ID", "Objective", "Pilot success measure"], DATA["objectives"], [0.55, 2.2, 3.75], font_size=9.5)

doc.add_heading("MVP scope", level=2)
for item in ["Single-owner account, onboarding and practice settings", "Clients, contacts and balances", "Matters, court details, tasks and financial summaries", "Web/manual/Gmail/research work capture with recovery", "Owner-operated work review: Draft -> Ready to Bill -> Excluded", "Rates, billables, expenses and unbilled WIP", "Professional GST invoice, PDF, delivery and revisions", "Full/partial payments, receipts, write-offs and aging", "Daily dashboard, income, WIP and outstanding reports"]:
    add_bullet(doc, item)
doc.add_heading("Explicit non-goals", level=2)
for item in DATA["nonGoals"]:
    add_bullet(doc, item)

h = doc.add_heading("2. Scrum delivery model", level=1)
h.paragraph_format.page_break_before = True
doc.add_paragraph("The implementation uses four two-week sprints. The solo developer is Product Owner, engineer and release owner; AI is used for bounded implementation and test tasks, while product decisions, integration, security and financial verification remain human-owned.")
add_table(doc, ["Sprint", "Goal", "Branches", "Plan", "Buffer", "Sprint review demo"], [
    [s["id"] + "\n" + s["weeks"], s["goal"], s["branches"], str(s["plannedHours"]) + " h", str(s["bufferHours"]) + " h", s["demo"]] for s in DATA["sprints"]
], [0.8, 1.85, 0.55, 0.55, 0.55, 2.2], font_size=8.5)
doc.add_heading("Cadence", level=2)
for text in ["Sprint planning: first morning of each sprint; confirm sprint goal, branch sequence and acceptance tests.", "Daily planning: 15 minutes; select one shippable vertical slice and state the verification command before coding.", "Daily review: 15 minutes; update tracker status, hours remaining, blockers and decision log.", "Backlog refinement: 45 minutes each Friday; split oversized items and protect the frozen non-goals.", "Sprint review: working demonstration against real API data and release criteria.", "Retrospective: record one process improvement and apply it in the next sprint."]:
    add_bullet(doc, text)
doc.add_heading("Definition of Ready", level=2)
for text in ["User outcome and route are named.", "Relevant frontend and backend files are identified.", "Data ownership and workspace-scope rules are explicit.", "Acceptance criteria are testable.", "Dependencies and migration impact are known."]:
    add_bullet(doc, text)
doc.add_heading("Definition of Done", level=2)
for text in ["Implementation is connected to real APIs; no hidden placeholder remains.", "Success, loading, empty, validation, error and retry states work.", "Workspace ownership and input validation are covered by tests.", "Frontend build and relevant backend/extension tests pass.", "The branch tracker, API mapping and known limitations are updated.", "The vertical workflow is manually verified on desktop and mobile where relevant."]:
    add_bullet(doc, text)

doc.add_heading("Sprint exit criteria", level=2)
add_table(doc, ["Sprint", "Exit criteria"], [[s["id"], s["exit"]] for s in DATA["sprints"]], [0.75, 5.75], font_size=9.5)

h = doc.add_heading("3. Release gates and governance", level=1)
h.paragraph_format.page_break_before = True
for idx, gate in enumerate(DATA["releaseGates"], start=1):
    add_number(doc, gate)
doc.add_heading("Top delivery risks", level=2)
add_table(doc, ["ID", "Severity", "Risk", "Mitigation", "Due"], [[r[0],r[1],r[2],r[4],r[5]] for r in DATA["risks"]], [0.45,0.7,1.7,3.05,0.6], font_size=8.5)
doc.add_heading("Change-control rule", level=2)
add_callout(doc, "Scope protection", "No new product module enters the MVP after Sprint 1. A new request must replace equal or greater planned work, preserve release gates, and be recorded in the decision log.", fill="FFF8E6", accent=GOLD)
h = doc.add_heading("4. Branch and work-package index", level=1)
h.paragraph_format.page_break_before = True
add_table(doc, ["No.", "Sprint", "Branch", "Epic", "Hours", "SP", "Primary outcome"], [[b["no"],b["sprint"],b["branch"],b["epic"],b["hours"],b["points"],b["goal"]] for b in DATA["branches"]], [0.35,0.45,1.65,0.8,0.45,0.35,2.45], font_size=7.8)

for b in DATA["branches"]:
    h = doc.add_heading(f'{b["no"]} - {b["branch"]}', level=1)
    h.paragraph_format.page_break_before = True
    p = doc.add_paragraph()
    r = p.add_run(f'{b["sprint"]}  |  {b["epic"]}  |  {b["hours"]} planned hours  |  {b["points"]} story points')
    set_font(r, size=10, bold=True, color=TEAL)
    add_callout(doc, "Outcome", b["goal"])
    add_table(doc, ["Delivery field", "Branch requirement"], [
        ["Branch", b["branch"]],
        ["Screens / states", b["screens"]],
        ["Backend / API", b["backend"]],
        ["Acceptance criteria", b["acceptance"]],
        ["Verification", b["verify"]],
        ["Commit", b["commit"]]
    ], [1.4,5.1], font_size=9.5)
    doc.add_heading("Codex implementation prompt", level=2)
    prompt = (
        "You are a senior full-stack engineer implementing Lexora Solo, a focused legal work-to-payment product for independent lawyers. "
        f"Work on branch `{b['branch']}` based on `codex/solo-lawyer-mvp`. The required outcome is: {b['goal']} "
        f"Implement these screens and states: {b['screens']}. Backend/API scope: {b['backend']}. "
        "Inspect existing routes, controllers, models, validators, tests and frontend API clients before editing. Reuse existing components and preserve legacy data compatibility. "
        "All retained records must be scoped to the authenticated workspace; never trust ownership fields from the client. Use real API data and honest recovery states. "
        f"The work is complete only when: {b['acceptance']}"
    )
    p = doc.add_paragraph(prompt)
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.right_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs: set_font(run, size=10.2, color=INK)
    doc.add_heading("Implementation guardrails", level=2)
    guards = [
        "Make additive schema changes first; migrate and verify before removing legacy fields or states.",
        "Prefer one canonical workflow over parallel conversion paths.",
        "Calculate financial totals server-side and preserve idempotency for every conversion.",
        "Add tests with the implementation, including a failure or boundary case.",
        "Do not reintroduce multi-user approvals, RBAC, HR/payroll or enterprise administration.",
        "Do not commit secrets, generated build output or unrelated user changes."
    ]
    for g in guards: add_bullet(doc, g)
    doc.add_heading("Required completion response", level=2)
    for response_no, item in enumerate(["Implementation summary", "Screens/routes completed", "API/schema changes", "Migration or compatibility notes", "Tests and exact verification commands", "Known limitations", f"Suggested commit: {b['commit']}"], start=1):
        add_manual_number(doc, response_no, item)

h = doc.add_heading("5. Screen inventory", level=1)
h.paragraph_format.page_break_before = True
add_table(doc, ["Module", "Screen", "Route", "Action", "Sprint", "Branch"], DATA["screens"], [0.75,1.55,1.65,0.9,0.55,0.45], font_size=8.3)

doc.add_heading("6. Pilot launch playbook", level=1)
doc.add_heading("Pilot cohort", level=2)
doc.add_paragraph("Recruit 3-5 independent lawyers who currently track work manually or in spreadsheets. Use separate workspaces, realistic client/matter data and at least one invoice cycle per participant.")
doc.add_heading("Pilot week measures", level=2)
for item in ["Activation: account created, practice settings completed, first client and matter created.", "Capture: first timer/manual/Gmail entry recorded and reviewed.", "Billing: first ready-to-bill item and invoice generated.", "Collection: payment recorded or outstanding invoice correctly aged.", "Quality: capture failures, duplicate prevention events, invoice corrections and support requests."]:
    add_bullet(doc, item)
doc.add_heading("Go / no-go decision", level=2)
add_callout(doc, "Go", "Proceed beyond controlled pilot only when all release gates pass, at least three lawyers complete the end-to-end flow, and no financial or isolation defect remains unresolved.", fill="E9F7F5", accent=GREEN)

doc.add_heading("7. Future integration boundary", level=1)
doc.add_paragraph("After the solo workflow is stable, expose bounded integration contracts instead of expanding into a full ERP. Clients and matters may sync from CRM/practice systems; invoices and payments may sync with accounting systems.")
doc.add_paragraph("Prioritised options: workspace-scoped REST API; outbound webhooks for work-ready, invoice-sent and payment-cleared events; CSV migration/accountant handoff; Zoho CRM/Books, QuickBooks or Xero adapters; Google/Microsoft calendar connections; and practice-management platform connectors.")

output = OUT / "Lexora_Solo_Lawyer_MVP_Product_and_Scrum_Plan.docx"
doc.save(output)
print(output)
